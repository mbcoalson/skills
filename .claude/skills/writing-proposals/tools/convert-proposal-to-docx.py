#!/usr/bin/env python3
"""
Convert Proposal Markdown to Word (.docx) with Template Cover Pages

This specialized tool converts proposal markdown files to Word documents while:
1. Preserving template cover pages (pages 1-2) with images and formatting
2. Applying template styling to proposal content (headers, body, tables)
3. Ensuring consistent formatting throughout the document

The tool is specifically designed for Icon Energy Consulting proposal workflow
where the proposal-template.docx contains:
- Page 1: Cover page with company branding and images
- Page 2: Additional branding/introduction page
- Page 3+: Proposal content begins (styled according to template)

Requirements:
- pypandoc: pip install pypandoc
- python-docx: pip install python-docx
- Pandoc will be auto-downloaded by pypandoc if not installed

Usage:
    python convert-proposal-to-docx.py input.md
    python convert-proposal-to-docx.py input.md output.docx
    python convert-proposal-to-docx.py input.md --template custom-template.docx
"""

import sys
import os
import json
import argparse
import subprocess
from pathlib import Path

try:
    import pypandoc
except ImportError:
    print("Error: pypandoc is not installed")
    print("Install with: pip install pypandoc")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)


def setup_pandoc():
    """
    Ensure Pandoc is installed. If not, download it automatically.
    """
    try:
        version = pypandoc.get_pandoc_version()
        print(f"[OK] Pandoc version {version} detected")
        return True
    except OSError:
        print("Pandoc not found. Downloading Pandoc...")
        try:
            pypandoc.download_pandoc()
            print("[OK] Pandoc downloaded successfully")
            return True
        except Exception as e:
            print(f"[ERROR] Error downloading Pandoc: {e}")
            print("\nManual installation:")
            print("  1. Visit: https://pandoc.org/installing.html")
            print("  2. Or download from: https://github.com/jgm/pandoc/releases")
            return False


def load_template_corrections(template_path):
    """
    Load learned corrections for the specified template.

    Returns:
        dict: Corrections data for this template, or empty dict if none found
    """
    script_dir = Path(__file__).parent
    corrections_file = script_dir / "template-corrections.json"

    if not corrections_file.exists():
        return {}

    try:
        with open(corrections_file, 'r') as f:
            db = json.load(f)

        template_name = Path(template_path).name
        return db.get(template_name, {})
    except Exception as e:
        print(f"[WARNING] Could not load template corrections: {e}")
        return {}


def apply_learned_corrections(doc, template_corrections):
    """
    Apply previously learned corrections to the document.

    Args:
        doc: python-docx Document object
        template_corrections: dict of corrections for this template

    Returns:
        dict: Statistics on applied corrections
    """
    stats = {
        "style_corrections": 0,
        "find_replace": 0
    }

    if not template_corrections:
        return stats

    # Apply additional style corrections beyond the built-in ones
    additional_styles = template_corrections.get("style_corrections", {})
    if additional_styles:
        print(f"  Applying {len(additional_styles)} learned style corrections...")

        for para in doc.paragraphs:
            if para.style.name in additional_styles:
                target_style = additional_styles[para.style.name]
                try:
                    para.style = target_style
                    stats["style_corrections"] += 1
                except:
                    pass

    # Apply find/replace corrections
    find_replace_rules = template_corrections.get("find_replace", [])
    if find_replace_rules:
        print(f"  Applying {len(find_replace_rules)} learned find/replace rules...")

        for rule in find_replace_rules:
            find_text = rule.get("find", "")
            replace_text = rule.get("replace", "")

            if not find_text:
                continue

            for para in doc.paragraphs:
                if find_text in para.text:
                    # Simple text replacement
                    para.text = para.text.replace(find_text, replace_text)
                    stats["find_replace"] += 1

    return stats


def convert_markdown_to_temp_docx(input_file, reference_doc):
    """
    Convert markdown to temporary docx using pypandoc with reference template.
    This applies all the template's styles to the converted content.

    Args:
        input_file (Path): Path to input markdown file
        reference_doc (Path): Path to reference template

    Returns:
        Path: Path to temporary converted file, or None if conversion failed
    """
    temp_output = input_file.parent / f"{input_file.stem}_temp_converted.docx"

    print(f"\n[Step 1/3] Converting markdown to styled document...")
    print(f"  Input: {input_file.name}")
    print(f"  Reference template: {reference_doc.name}")

    try:
        # Convert using pypandoc with reference doc for styling
        pypandoc.convert_file(
            str(input_file),
            'docx',
            outputfile=str(temp_output),
            extra_args=[f'--reference-doc={reference_doc}']
        )
        print(f"[OK] Styled conversion complete")
        return temp_output

    except RuntimeError as e:
        print(f"[ERROR] Pandoc conversion error: {e}")
        return None
    except Exception as e:
        print(f"[ERROR] Unexpected error during conversion: {e}")
        return None


def merge_template_and_content(template_path, converted_path, output_path):
    """
    Merge template cover pages (1-2) with converted proposal content.

    Strategy:
    1. Make a copy of the template file to work with
    2. Open the copy (preserves all formatting and layout)
    3. Find "Building Information" marker
    4. Remove all elements from that marker onward
    5. Load converted proposal content
    6. Append proposal elements to the template copy
    7. Apply built-in style corrections (Normal -> Body Text, etc.)
    8. Apply learned corrections from previous feedback
    9. Save as final output

    This approach preserves template pages 1-2 by starting with an actual copy
    of the template file, ensuring all layout relationships stay intact.

    Args:
        template_path (Path): Path to template with cover pages
        converted_path (Path): Path to converted proposal content
        output_path (Path): Path for final merged output

    Returns:
        bool: True if successful, False otherwise
    """
    print(f"\n[Step 2/3] Merging template cover pages with proposal content...")

    # Load learned corrections for this template
    template_corrections = load_template_corrections(template_path)
    if template_corrections:
        automated_count = sum(1 for f in template_corrections.get("formatting_fixes", []) if f.get("automated"))
        pending_count = sum(1 for f in template_corrections.get("formatting_fixes", []) if not f.get("automated"))
        print(f"  Found {automated_count} automated corrections for this template")
        if pending_count > 0:
            print(f"  ({pending_count} corrections pending implementation)")

    try:
        import shutil

        # Create a temporary copy of the template to work with
        temp_template = output_path.parent / f"temp_template_{output_path.stem}.docx"
        shutil.copy2(template_path, temp_template)
        print(f"  Created temporary template copy")

        # Open the template copy (this preserves all formatting perfectly)
        final_doc = Document(str(temp_template))
        print(f"  Template copy loaded: {len(final_doc.paragraphs)} paragraphs, {len(final_doc.sections)} sections")

        # Find where page 3 starts by looking for "Building Information" header
        marker_element = None
        for para in final_doc.paragraphs:
            if para.text.strip() == "Building Information":
                marker_element = para._element
                print(f"  Found 'Building Information' marker")
                break

        if marker_element is None:
            print(f"  [WARNING] 'Building Information' marker not found")
            elements_to_remove = []
        else:
            # Get all body elements
            body_elements = list(final_doc.element.body)

            # Find the index of the marker element
            marker_index = None
            for idx, elem in enumerate(body_elements):
                if elem == marker_element:
                    marker_index = idx
                    break

            if marker_index is not None:
                # Get all elements from marker onward (including marker itself)
                elements_to_remove = body_elements[marker_index:]
                print(f"  Found {len(elements_to_remove)} elements to remove (from 'Building Information' onward)")
            else:
                print(f"  [WARNING] Could not find marker in body elements")
                elements_to_remove = []

        # Remove all elements from the marker onward
        removed_count = 0
        for element in elements_to_remove:
            try:
                element.getparent().remove(element)
                removed_count += 1
            except:
                pass

        print(f"  Removed {removed_count} elements after 'Building Information' marker")

        # Load converted proposal content
        content_doc = Document(str(converted_path))
        print(f"  Proposal content loaded: {len(content_doc.paragraphs)} paragraphs, {len(content_doc.tables)} tables")

        # Append all proposal content elements to the template copy
        for element in content_doc.element.body:
            final_doc.element.body.append(element)

        print(f"  Appended proposal content to template")

        # Clean up the temporary template file
        try:
            temp_template.unlink()
        except:
            pass

        # Remap styles to template defaults
        # Find template styles
        body_text_style = None
        body_list_style = None
        list_paragraph_style = None

        for style in final_doc.styles:
            if style.name == 'Body Text':
                body_text_style = style
            elif style.name == 'Body List':
                body_list_style = style
            elif style.name == 'List Paragraph':
                list_paragraph_style = style

        if body_text_style:
            normal_to_body = 0
            bullets_remapped = 0
            skip_count = 0

            for para in final_doc.paragraphs:
                # Skip headings and titles (these should stay as-is)
                if (para.style.name.startswith('Heading') or
                    para.style.name in ['Title', 'Subtitle']):
                    skip_count += 1
                    continue

                # Only process paragraphs with text
                if not para.text.strip():
                    continue

                try:
                    # Check if this is a bullet/list paragraph
                    is_list = False
                    if para._element.pPr is not None:
                        numPr = para._element.pPr.numPr
                        if numPr is not None:
                            is_list = True

                    if is_list:
                        # Remap list paragraphs to Body List style if available
                        if body_list_style:
                            para.style = body_list_style
                            bullets_remapped += 1
                        elif list_paragraph_style:
                            para.style = list_paragraph_style
                            bullets_remapped += 1
                    else:
                        # Remap Normal and other body styles to Body Text
                        if para.style.name in ['Normal', 'Body Text', 'First Paragraph']:
                            para.style = body_text_style
                            if para.style.name == 'Normal':
                                normal_to_body += 1
                except:
                    pass  # Skip if style can't be applied

            print(f"  Style remapping complete:")
            print(f"    - Normal -> Body Text: {normal_to_body} paragraphs")
            print(f"    - Bullets remapped: {bullets_remapped} items")
            print(f"    - Headings/titles skipped: {skip_count}")
        else:
            print(f"  [WARNING] Body Text style not found in template")

        # Apply learned corrections from previous feedback
        if template_corrections:
            correction_stats = apply_learned_corrections(final_doc, template_corrections)
            if correction_stats["style_corrections"] > 0 or correction_stats["find_replace"] > 0:
                print(f"  Learned corrections applied:")
                if correction_stats["style_corrections"] > 0:
                    print(f"    - Style corrections: {correction_stats['style_corrections']}")
                if correction_stats["find_replace"] > 0:
                    print(f"    - Find/replace: {correction_stats['find_replace']}")

        # Save the merged document
        final_doc.save(str(output_path))
        print(f"[OK] Documents merged successfully")
        return True

    except Exception as e:
        print(f"[ERROR] Error merging documents: {e}")
        import traceback
        traceback.print_exc()
        return False


def cleanup_temp_file(temp_file):
    """Remove temporary conversion file."""
    try:
        if temp_file and temp_file.exists():
            temp_file.unlink()
            print(f"[OK] Temporary file cleaned up")
    except Exception as e:
        print(f"[WARNING] Could not remove temporary file: {e}")


def convert_proposal(input_file, output_file=None, template_path=None, collect_feedback=False):
    """
    Main conversion workflow.

    Args:
        input_file (str): Path to input markdown file
        output_file (str, optional): Path to output Word file
        template_path (str, optional): Path to template file
        collect_feedback (bool): Whether to collect feedback after conversion

    Returns:
        tuple: (success: bool, output_path: Path or None)
    """
    # Validate input file
    input_path = Path(input_file)
    if not input_path.exists():
        print(f"[ERROR] Input file not found: {input_file}")
        return False, None

    # Determine template path
    if template_path is None:
        # Default to writing-proposals skill template
        script_dir = Path(__file__).parent
        template_path = script_dir.parent / "templates" / "proposal-template.docx"
    else:
        template_path = Path(template_path)

    if not template_path.exists():
        print(f"[ERROR] Template file not found: {template_path}")
        return False, None

    # Determine output file
    if output_file is None:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(output_file)
        output_path.parent.mkdir(parents=True, exist_ok=True)

    print("=" * 70)
    print("PROPOSAL CONVERSION TO WORD")
    print("=" * 70)
    print(f"Input file: {input_path}")
    print(f"Template: {template_path}")
    print(f"Output file: {output_path}")
    print("=" * 70)

    # Step 1: Convert markdown to styled docx (temporary)
    temp_converted = convert_markdown_to_temp_docx(input_path, template_path)
    if not temp_converted:
        return False, None

    # Step 2: Merge template cover pages with converted content
    success = merge_template_and_content(template_path, temp_converted, output_path)

    # Step 3: Cleanup
    print(f"\n[Step 3/3] Cleanup...")
    cleanup_temp_file(temp_converted)

    if success:
        print("\n" + "=" * 70)
        print("[SUCCESS] CONVERSION COMPLETE")
        print("=" * 70)
        print(f"Output: {output_path.absolute()}")
        print(f"Template: {template_path.name}")
        print("\nNext steps:")
        print("  1. Open the output document in Word")
        print("  2. Verify pages 1-2 match template (cover pages)")
        print("  3. Verify page 3+ formatting matches template styles")
        print("  4. Review all tables, headers, and body text formatting")
        print("  5. Make any final manual adjustments as needed")

        if collect_feedback:
            print("\n  >> FEEDBACK COLLECTION ENABLED")
            print("     After reviewing the document, you'll be asked about")
            print("     manual corrections needed. This helps improve future")
            print("     conversions using this template.")
        else:
            print("\nTo help improve future conversions, run:")
            print(f"  python tools/collect-feedback.py \"{output_path}\" --template \"{template_path.name}\"")

        print("=" * 70)
        return True, output_path
    else:
        print("\n[ERROR] Conversion failed")
        return False, None


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description='Convert proposal markdown to Word with template cover pages',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s proposal.md
  %(prog)s proposal.md final-proposal.docx
  %(prog)s proposal.md --template custom-template.docx

Template Requirements:
  - Pages 1-2: Cover pages with branding/images (preserved exactly)
  - Styles: Headers, body text, table styles (applied to proposal content)

Default template location:
  .claude/skills/writing-proposals/templates/proposal-template.docx
        """
    )

    parser.add_argument(
        'input_file',
        help='Input markdown file (.md)'
    )

    parser.add_argument(
        'output_file',
        nargs='?',
        default=None,
        help='Output Word file (.docx) - defaults to same name as input'
    )

    parser.add_argument(
        '--template', '-t',
        dest='template_path',
        help='Custom template file (default: skill template)'
    )

    parser.add_argument(
        '--feedback', '-f',
        action='store_true',
        help='Collect feedback on manual corrections after conversion'
    )

    parser.add_argument(
        '--version',
        action='version',
        version='%(prog)s 2.0.0'
    )

    args = parser.parse_args()

    # Ensure Pandoc is installed
    if not setup_pandoc():
        sys.exit(1)

    # Perform conversion
    success, output_path = convert_proposal(
        args.input_file,
        args.output_file,
        args.template_path,
        args.feedback
    )

    # Collect feedback if requested and conversion succeeded
    if success and args.feedback and output_path:
        print("\n" + "=" * 70)
        print("STARTING FEEDBACK COLLECTION")
        print("=" * 70)
        print("\nPlease review the Word document, then return here to provide feedback.")
        input("\nPress Enter when you're ready to provide feedback...")

        # Call collect-feedback.py
        script_dir = Path(__file__).parent
        feedback_script = script_dir / "collect-feedback.py"

        feedback_args = [sys.executable, str(feedback_script), str(output_path)]
        if args.template_path:
            feedback_args.extend(['--template', args.template_path])

        try:
            subprocess.run(feedback_args)
        except Exception as e:
            print(f"[ERROR] Could not run feedback collection: {e}")
            print(f"\nYou can run it manually:")
            print(f"  python tools/collect-feedback.py \"{output_path}\"")

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
