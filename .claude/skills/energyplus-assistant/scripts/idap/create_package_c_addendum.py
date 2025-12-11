#!/usr/bin/env python3
"""
Create IDAP SD Report Addendum - Package C Updated Model
Generates Word document describing updated Package C system configuration
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from datetime import datetime

def extract_table_by_heading(soup, heading_text):
    """Extract HTML table following a specific heading"""
    for heading in soup.find_all(['h2', 'h3']):
        if heading_text.lower() in heading.text.lower():
            table = heading.find_next('table')
            if table:
                rows = []
                for tr in table.find_all('tr'):
                    cells = [td.text.strip() for td in tr.find_all(['td', 'th'])]
                    if cells:
                        rows.append(cells)
                return rows
    return []

def add_formatted_heading(doc, text, level=1):
    """Add heading with consistent formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table(doc, headers, data, style='Light Grid Accent 1'):
    """Add formatted table"""
    if not data:
        doc.add_paragraph('[No data available]')
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style

    # Headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data[:len(headers)]):  # Prevent index errors
            row_cells[i].text = str(cell_data)

    doc.add_paragraph()  # Spacing

def main():
    if len(sys.argv) < 2:
        print("Usage: python create_package_c_addendum.py <openstudio_results.html> [output.docx]")
        sys.exit(1)

    html_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "SECC_Package_C_Addendum.docx"

    if not os.path.exists(html_path):
        print(f"[ERROR] HTML file not found: {html_path}")
        sys.exit(1)

    print("=" * 80)
    print("PACKAGE C ADDENDUM GENERATOR")
    print("=" * 80)
    print(f"Source: {html_path}")
    print(f"Output: {output_path}\n")

    # Load HTML
    print("Loading OpenStudio results...")
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    # Extract data
    print("Extracting equipment data...")
    airloop_data = extract_table_by_heading(soup, "Air Loops Detail")
    plant_data = extract_table_by_heading(soup, "Plant Loops Detail")
    zone_equip_data = extract_table_by_heading(soup, "Zone Equipment Detail")

    # Create document
    print("Creating Word document...")
    doc = Document()

    # Title page
    title = doc.add_heading('Southeast Community Center', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Schematic Design Energy Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Addendum: Updated Package C Model', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Project info
    p = doc.add_paragraph()
    p.add_run('IDAP Project # PR-600523\n').bold = True
    p.add_run(f'{datetime.now().strftime("%B %d, %Y")}\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Purpose
    add_formatted_heading(doc, 'Purpose of This Addendum', level=1)

    doc.add_paragraph(
        'This addendum updates Package C (Ground Source Heat Pump System) design details '
        'based on the latest energy model incorporating Engineer of Record (EOR) HVAC specifications. '
        'The updated model includes:'
    )

    updates = doc.add_paragraph(style='List Bullet')
    updates.add_run('Water-to-air heat pump rooftop units (HP RTU) with water-side economizer\n')
    updates = doc.add_paragraph(style='List Bullet')
    updates.add_run('Dedicated outdoor air systems (HP DOAS) with energy recovery wheels\n')
    updates = doc.add_paragraph(style='List Bullet')
    updates.add_run('Zone-level water-source heat pumps for local temperature control\n')
    updates = doc.add_paragraph(style='List Bullet')
    updates.add_run('Condenser water loop served by two-speed evaporative fluid cooler\n')
    updates = doc.add_paragraph(style='List Bullet')
    updates.add_run('Hot water loop with condensing boilers for backup heating\n')
    updates = doc.add_paragraph(style='List Bullet')
    updates.add_run('Pool dehumidification system with dedicated controls\n')

    doc.add_paragraph()
    doc.add_paragraph(
        'Energy performance results and IDAP incentive calculations will be provided '
        'upon completion of the baseline model comparison.'
    )

    doc.add_page_break()

    # Updated Package C Description
    add_formatted_heading(doc, 'Updated Package C: Ground Source Heat Pump System Description', level=1)

    add_formatted_heading(doc, 'System Overview', level=2)

    doc.add_paragraph(
        'The Package C design utilizes a hybrid ground-source heat pump system optimized for the '
        'SECC facility mix of recreation, library, fitness, and aquatics spaces. The system provides:'
    )

    doc.add_paragraph('Primary heating and cooling via water-to-air heat pump equipment', style='List Bullet')
    doc.add_paragraph('Energy recovery on all outdoor air systems', style='List Bullet')
    doc.add_paragraph('Condensing boiler backup for peak heating and pool loads', style='List Bullet')
    doc.add_paragraph('Two-speed evaporative fluid cooler for heat rejection and water-side economizer operation', style='List Bullet')

    doc.add_page_break()

    # Air Systems
    add_formatted_heading(doc, 'Air Distribution Systems', level=2)

    doc.add_paragraph(
        'The facility is served by 7 air handling systems, comprising 5 heat pump rooftop units (HP RTU), '
        '2 heat pump dedicated outdoor air systems (HP DOAS), and 1 pool dehumidification unit.'
    )
    doc.add_paragraph()

    # Extract airloop names from table
    if airloop_data and len(airloop_data) > 1:
        airloop_names = []
        for row in airloop_data[1:]:  # Skip header
            if row and len(row) > 0:
                airloop_names.append(row[0])

        if airloop_names:
            add_formatted_heading(doc, 'Air System Configuration', level=3)

            # Group by type
            doas_systems = [name for name in airloop_names if 'DOAS' in name]
            rtu_systems = [name for name in airloop_names if 'RTU' in name]
            pool_systems = [name for name in airloop_names if 'Pool' in name]

            if doas_systems:
                doc.add_paragraph().add_run('Heat Pump DOAS Units:').bold = True
                for name in doas_systems:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{name}\n')
                    p.add_run('  • Water-to-air heat pump with water-side economizer\n')
                    p.add_run('  • Rotary energy recovery wheel\n')
                    p.add_run('  • Electric preheat coil\n')
                    p.add_run('  • Provides ventilation air to VAV terminals serving zones\n')

            if rtu_systems:
                doc.add_paragraph().add_run('Heat Pump Rooftop Units:').bold = True
                for name in rtu_systems:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{name}\n')
                    p.add_run('  • Water-to-air heat pump with water-side economizer\n')
                    p.add_run('  • Rotary energy recovery wheel\n')
                    p.add_run('  • Electric preheat coil\n')
                    p.add_run('  • VAV distribution with zone-level equipment\n')

            if pool_systems:
                doc.add_paragraph().add_run('Pool Dehumidification System:').bold = True
                for name in pool_systems:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{name}\n')
                    p.add_run('  • DX cooling with humidity control\n')
                    p.add_run('  • Hot water heating coil\n')
                    p.add_run('  • Rotary energy recovery wheel\n')
                    p.add_run('  • Electric reheat for precise dewpoint control\n')

    doc.add_page_break()

    # Central Plant
    add_formatted_heading(doc, 'Central Plant Systems', level=2)

    add_formatted_heading(doc, 'Condenser Water Loop', level=3)
    doc.add_paragraph(
        'A central condenser water loop serves the water-side economizer coils in all HP RTU and HP DOAS units. '
        'This loop is cooled by a two-speed evaporative fluid cooler, enabling efficient heat rejection '
        'and enabling water-side economizer operation during favorable outdoor conditions.'
    )
    doc.add_paragraph()

    doc.add_paragraph().add_run('Key Features:').bold = True
    doc.add_paragraph('Two-speed evaporative fluid cooler for heat rejection', style='List Bullet')
    doc.add_paragraph('Variable primary pumping with differential pressure reset', style='List Bullet')
    doc.add_paragraph('Condenser water supply temperature reset based on outdoor wet-bulb', style='List Bullet')
    doc.add_paragraph('Integration with water-side economizer coils in air systems', style='List Bullet')

    doc.add_paragraph()

    add_formatted_heading(doc, 'Hot Water Loop', level=3)
    doc.add_paragraph(
        'A hot water loop provides backup heating to air system coils and serves pool heating loads. '
        'The loop is served by condensing natural gas boilers staged to meet demand efficiently.'
    )
    doc.add_paragraph()

    doc.add_paragraph().add_run('Key Features:').bold = True
    doc.add_paragraph('Four condensing hot water boilers (natural gas, 80% thermal efficiency)', style='List Bullet')
    doc.add_paragraph('Primary-only pumping arrangement with variable speed drives', style='List Bullet')
    doc.add_paragraph('Hot water supply temperature reset based on outdoor air temperature', style='List Bullet')
    doc.add_paragraph('Sequenced boiler staging to optimize part-load efficiency', style='List Bullet')

    doc.add_page_break()

    # Zone Equipment
    add_formatted_heading(doc, 'Zone-Level Equipment', level=2)

    doc.add_paragraph(
        'Individual thermal zones are served by water-source heat pumps connected to the condenser water loop, '
        'providing local temperature control and load diversity benefits.'
    )
    doc.add_paragraph()

    if zone_equip_data and len(zone_equip_data) > 1:
        # Count WSHPs
        wshp_count = sum(1 for row in zone_equip_data[1:] if row and 'Heat Pump' in str(row))

        doc.add_paragraph(f'Total zone-level water-source heat pumps: {wshp_count} units')
        doc.add_paragraph()

    doc.add_paragraph().add_run('Benefits:').bold = True
    doc.add_paragraph('Simultaneous heating and cooling capability with heat recovery through condenser loop', style='List Bullet')
    doc.add_paragraph('Individual zone control for diverse space types', style='List Bullet')
    doc.add_paragraph('Reduced distribution losses compared to central air systems', style='List Bullet')
    doc.add_paragraph('Load diversity reduces peak plant capacity requirements', style='List Bullet')

    doc.add_page_break()

    # Control Strategies
    add_formatted_heading(doc, 'Control Strategies and Sequences', level=1)

    add_formatted_heading(doc, 'DOAS Control', level=2)
    doc.add_paragraph('Supply-air dewpoint control (48-52°F) to manage latent loads', style='List Bullet')
    doc.add_paragraph('Energy recovery wheel operation integrated with outdoor air economizer', style='List Bullet')
    doc.add_paragraph('Static pressure reset based on zone damper positions', style='List Bullet')
    doc.add_paragraph('Demand-controlled ventilation tracking occupancy', style='List Bullet')

    doc.add_paragraph()

    add_formatted_heading(doc, 'Zone WSHP Control', level=2)
    doc.add_paragraph('Local zone temperature control with occupied/unoccupied setbacks', style='List Bullet')
    doc.add_paragraph('Loop temperature limits to protect compressor operation', style='List Bullet')
    doc.add_paragraph('Coordination with DOAS for dehumidification priority', style='List Bullet')

    doc.add_paragraph()

    add_formatted_heading(doc, 'Condenser Water Loop Control', level=2)
    doc.add_paragraph('Loop water temperature maintained within 65-85°F operating band', style='List Bullet')
    doc.add_paragraph('Condenser water supply temperature reset based on outdoor wet-bulb + 3°F', style='List Bullet')
    doc.add_paragraph('Two-speed fluid cooler staging based on loop temperature', style='List Bullet')
    doc.add_paragraph('Variable primary pumping with differential pressure reset', style='List Bullet')
    doc.add_paragraph('Water-side economizer enable when outdoor conditions favorable', style='List Bullet')

    doc.add_paragraph()

    add_formatted_heading(doc, 'Hot Water Loop Control', level=2)
    doc.add_paragraph('Hot water supply temperature reset from 130°F down to 100°F based on outdoor air', style='List Bullet')
    doc.add_paragraph('Boiler staging: lead boiler carries base load, lag boilers stage on as needed', style='List Bullet')
    doc.add_paragraph('Boiler lockout when loop temperature can be maintained by other sources', style='List Bullet')
    doc.add_paragraph('Priority sequencing: pool loads, then space heating, minimize simultaneous operation', style='List Bullet')

    doc.add_page_break()

    # Modeling Results Placeholder
    add_formatted_heading(doc, 'Modeling Results (Pending Baseline Comparison)', level=1)

    doc.add_paragraph(
        'The following results will be populated upon completion of the ASHRAE 90.1-2019 Appendix G '
        'baseline model comparison:'
    )
    doc.add_paragraph()

    # Table 1 structure
    add_formatted_heading(doc, 'Table 1: Updated Package C Modeling Results', level=2)

    results_data = [
        ['Design Package', 'Package C: Ground Source Heat Pump (Updated Model)'],
        ['Proposed Building Energy Cost (PBPm)', '[TBD - pending simulation]'],
        ['Energy Cost Savings vs. Code Baseline', '[TBD - requires BBPCode]'],
        ['Percent Below Code (%)', '[TBD - formula: 1 - PBREC/(BPFCode × BBREC)]'],
        ['Life Cycle Cost NPV ($)', '[TBD - requires first cost estimate]'],
        ['Peak Design Cooling Load (tons)', '[TBD - from model sizing]'],
        ['Peak Design Heating Load (kBtu/h)', '[TBD - from model sizing]'],
    ]

    add_table(doc, ['Metric', 'Value'], results_data)

    doc.add_paragraph(
        'Note: Calculation is based on regulated costs only per IDAP program requirements. '
        'Peak loads represent coil loads, not final equipment sizes, which must account for '
        'equipment efficiencies and safety factors per engineering standards.'
    )

    doc.add_page_break()

    # Construction Incentive
    add_formatted_heading(doc, 'Table 2: Updated Construction Incentive Estimate', level=2)

    ci_data = [
        ['Package', 'Package C: GSHP (Updated)'],
        ['PBPm ($/yr)', '[TBD]'],
        ['Construction Incentive ($)', '[TBD - formula: 2 × (BBPCode - PBPm) if PBPm ≤ PBPt]'],
    ]

    add_table(doc, ['Metric', 'Value'], ci_data)

    doc.add_paragraph(
        'Construction Incentive eligibility requires PBPm ≤ PBPt. '
        'If eligible, incentive equals 2× the annual energy cost savings vs. code baseline. '
        'Rule of thumb: every $10k/yr reduction in regulated cost increases CI by $20k.'
    )

    doc.add_page_break()

    # Performance Incentive
    add_formatted_heading(doc, 'Table 3: Updated Performance Incentive Estimate', level=2)

    pi_data = [
        ['Package', 'Package C: GSHP (Updated)'],
        ['PBREC ($/yr)*', '[TBD - modeled regulated energy cost]'],
        ['Performance Incentive ($)', '[TBD - formula: (BPFCode × BBREC) - Actual Regulated Cost]'],
    ]

    add_table(doc, ['Metric', 'Value'], pi_data)

    doc.add_paragraph('*Modeled PBREC shown as proxy for actual regulated energy cost')
    doc.add_paragraph()

    doc.add_paragraph(
        'Performance Incentive is based on actual sub-metered regulated energy cost over 12 consecutive '
        'months post-occupancy. Currently $0 (negative values floor to zero). With operational improvements '
        'and measured performance, every $10k/yr actual regulated cost falls below code baseline yields $10k in PI.'
    )

    doc.add_page_break()

    # Equipment Details from Model
    add_formatted_heading(doc, 'Detailed Equipment Schedule (As Modeled)', level=1)

    # Air Loops
    if airloop_data:
        add_formatted_heading(doc, 'Air Systems', level=2)

        # Filter and format headers
        if len(airloop_data) > 0:
            headers = airloop_data[0][:4]  # Take first 4 columns
            data = [row[:4] for row in airloop_data[1:] if row and len(row) > 0]

            if data:
                add_table(doc, headers, data)
            else:
                doc.add_paragraph(f'Total Air Systems: 7 (details in model)')

    # Plant Loops
    if plant_data:
        add_formatted_heading(doc, 'Plant Loops', level=2)

        if len(plant_data) > 0:
            headers = plant_data[0][:4]
            data = [row[:4] for row in plant_data[1:] if row and len(row) > 0]

            if data:
                add_table(doc, headers, data)

    # Zone Equipment
    if zone_equip_data:
        add_formatted_heading(doc, 'Zone Equipment Summary', level=2)

        doc.add_paragraph(
            'Individual zones are served by water-source heat pumps. '
            'Full equipment list available in detailed model documentation.'
        )
        doc.add_paragraph()

        # Show first few zones as example
        if len(zone_equip_data) > 1:
            headers = zone_equip_data[0][:3]
            data = [row[:3] for row in zone_equip_data[1:6] if row]  # First 5 zones

            if data:
                add_table(doc, headers, data)
                doc.add_paragraph(f'[... and {len(zone_equip_data) - 6} additional zones]')

    doc.add_page_break()

    # Next Steps
    add_formatted_heading(doc, 'Next Steps', level=1)

    doc.add_paragraph().add_run('To Complete IDAP SD Report for Updated Package C:').bold = True
    doc.add_paragraph()

    steps = [
        'Complete ASHRAE 90.1-2019 Appendix G baseline model',
        'Extract baseline performance metrics (BBPCode, BBREC)',
        'Calculate IDAP metrics for Package C updated model (PBPm, PBREC, % below code)',
        'Calculate Construction Incentive and Performance Incentive estimates',
        'Extract peak design loads from final model sizing',
        'Prepare life cycle cost analysis with first cost estimates',
        'Generate complete IDAP SD report with all three packages compared',
    ]

    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph(
        'Upon completion of these steps, this addendum will be integrated with the full '
        'IDAP Schematic Design Energy Report for submittal to Fort Collins Utilities IDAP program.'
    )

    # Save document
    doc.save(output_path)

    print(f"\n[OK] Addendum document created: {output_path}\n")
    print("Document contents:")
    print("  - Purpose and scope of addendum")
    print("  - Updated Package C system description")
    print("  - Air systems configuration (HP RTU, HP DOAS, Pool)")
    print("  - Central plant description (CW loop, HW loop)")
    print("  - Control strategies and sequences")
    print("  - Equipment schedule tables (extracted from OpenStudio)")
    print("  - Modeling results placeholders")
    print("  - Next steps for IDAP completion")
    print()
    print("[NOTE] Energy performance values marked [TBD] pending baseline comparison")

if __name__ == "__main__":
    main()
