#!/usr/bin/env python3
"""
Generate IDAP SD Report Addendum (Word Document)
Creates a Word document matching IDAP SD report format with equipment data
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from eppy.modeleditor import IDF

def find_idd():
    """Auto-detect Energy+.idd"""
    common_locations = [
        r'C:/EnergyPlusV25-1-0/Energy+.idd',
        r'C:/EnergyPlusV24-2-0/Energy+.idd',
    ]
    for loc in common_locations:
        if os.path.exists(loc):
            return loc
    return None

def add_formatted_heading(doc, text, level=1):
    """Add a heading with formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_equipment_table(doc, title, headers, data):
    """Add a formatted table to document"""
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    doc.add_paragraph()  # Spacing
    return table

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_idap_addendum.py <path_to_model.idf> [output.docx]")
        sys.exit(1)

    idf_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "SECC_IDAP_SD_Addendum.docx"

    if not os.path.exists(idf_path):
        print(f"[ERROR] IDF file not found: {idf_path}")
        sys.exit(1)

    # Find and load IDD
    idd_path = find_idd()
    if not idd_path:
        print("[ERROR] Could not find Energy+.idd file")
        sys.exit(1)

    print("=" * 80)
    print("IDAP SD REPORT ADDENDUM GENERATOR")
    print("=" * 80)
    print(f"IDF: {idf_path}")
    print(f"Output: {output_path}\n")

    # Load IDF
    try:
        IDF.setiddname(idd_path)
        idf = IDF(idf_path)
    except Exception as e:
        print(f"[ERROR] Failed to load IDF: {e}")
        sys.exit(1)

    print("Extracting equipment data...")

    # Extract equipment
    airloops = idf.idfobjects.get('AIRLOOPHVAC', [])
    boilers = idf.idfobjects.get('BOILER:HOTWATER', [])
    cooling_coils_wshp = idf.idfobjects.get('COIL:COOLING:WATERTOAIRHEATPUMP:EQUATIONFIT', [])
    heating_coils_wshp = idf.idfobjects.get('COIL:HEATING:WATERTOAIRHEATPUMP:EQUATIONFIT', [])
    fluid_coolers = idf.idfobjects.get('EVAPORATIVEFLUIDCOOLER:TWOSPEED', [])
    ervs = idf.idfobjects.get('HEATEXCHANGER:AIRTOAIR:SENSIBLEANDLATENT', [])
    var_pumps = idf.idfobjects.get('PUMP:VARIABLESPEED', [])
    const_pumps = idf.idfobjects.get('PUMP:CONSTANTSPEED', [])
    all_pumps = list(var_pumps) + list(const_pumps)

    # Create Word document
    doc = Document()

    # Title page
    title = doc.add_heading('Southeast Community Center', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Schematic Design Energy Report - Addendum', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('Package C: Ground Source Heat Pump System', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph(f'Model: SECC_WSHP_ProposedModel_v3')
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')

    doc.add_page_break()

    # Energy Model Input Details (Table 6 format)
    add_formatted_heading(doc, 'Energy Model Input Details', level=1)

    doc.add_paragraph('The following table summarizes the technical inputs to the Package C model as modeled.')
    doc.add_paragraph()

    # HVAC Primary Systems
    add_formatted_heading(doc, 'HVAC - Primary Systems (Central Plant)', level=2)

    primary_data = [
        ['Primary cooling system type', 'Ground-source water-to-air heat pumps on central ground loop; central heat pump plant serving DOAS CHW/HW coils'],
        ['Condenser type', 'Water-source tied to vertical borefield; supplemental evaporative fluid cooler, two speed'],
        ['Cooling Plant fuel type', 'Electric'],
        ['Primary Cooling efficiency', '[TBD from equipment sizing]'],
        ['CHW pumping arrangement and controls', '[TBD - extract from model]'],
        ['CHW pumps HP, GPM, ft head', '[TBD - extract from sizing]'],
        ['Primary heating system type', 'Water-to-air heat pumps on ground loop provide primary space heating; gas or electric boiler backup'],
        ['Heating plant fuel type', 'Natural gas'],
        ['Heating plant efficiency', '80% (gas boilers)'],
        ['HW pumping arrangement and controls', '[TBD - extract from model]'],
        ['HW pumps HP, GPM, ft head', '[TBD - extract from sizing]'],
    ]

    add_equipment_table(doc, '', ['Parameter', 'Value'], primary_data)

    # HVAC Secondary Systems
    add_formatted_heading(doc, 'HVAC - Secondary Systems', level=2)

    secondary_data = [
        ['Secondary system type', f'{len(airloops)} DOAS units with rotary ERV + HW/CHW coils; WSHPs for zone heating/cooling; Pool dehumidification unit'],
        ['Total supply fan BHP or kW & kW/CFM', '[TBD from equipment sizing]'],
        ['Total supply fan CFM & CFM/SF', '[TBD from equipment sizing]'],
        ['Supply fan control', 'DOAS fans VFD; WSHP indoor fans variable-speed'],
        ['Total return fan BHP or kW & kW/CFM', '[TBD - WSHP fan power]'],
        ['Total return fan CFM', '[TBD - delivered by WSHP fans]'],
        ['Return fan control', 'WSHP indoor fans variable-speed (ECM/VSD)'],
        ['Supply air temperature setpoint and reset', '[TBD - extract from schedules]'],
        ['Terminal unit reheat', 'N/A (zone WSHPs provide trim heat/cool; DOAS supplies ventilation/latent)'],
        ['Occupied Setpoints (Htg./Clg.)', '[TBD - extract from thermostats]'],
        ['Unoccupied Setpoints', '[TBD - extract from thermostats]'],
        ['Outside air control', 'DOAS with rotary energy recovery wheel; fixed minimum OA'],
    ]

    add_equipment_table(doc, '', ['Parameter', 'Value'], secondary_data)

    # Equipment Schedule
    doc.add_page_break()
    add_formatted_heading(doc, 'Equipment Schedule - Package C', level=1)

    doc.add_paragraph('The following equipment has been modeled in the SECC Package C (Ground Source Heat Pump) design:')
    doc.add_paragraph()

    # Air Systems
    add_formatted_heading(doc, 'Air Systems', level=2)

    airloop_data = []
    for i, airloop in enumerate(airloops, 1):
        airloop_data.append([
            str(i),
            airloop.Name,
            'DOAS' if 'DOAS' in airloop.Name else 'RTU',
            'Water-to-air heat pump with ERV'
        ])

    add_equipment_table(doc, '', ['#', 'System Name', 'Type', 'Configuration'], airloop_data)

    # Central Plant Equipment
    add_formatted_heading(doc, 'Central Plant Equipment', level=2)

    # Fluid Coolers
    if fluid_coolers:
        fc_data = []
        for fc in fluid_coolers:
            fc_data.append([
                fc.Name,
                'Two-Speed Evaporative Fluid Cooler',
                '[Capacity TBD from sizing]',
                'Two-speed fan control'
            ])
        add_equipment_table(doc, 'Evaporative Fluid Cooler', ['Name', 'Type', 'Capacity', 'Control'], fc_data)

    # Boilers
    if boilers:
        boiler_data = []
        for boiler in boilers:
            fuel = getattr(boiler, 'Fuel_Type', 'NaturalGas')
            eff = getattr(boiler, 'Nominal_Thermal_Efficiency', '0.80')
            boiler_data.append([
                boiler.Name,
                'Condensing Hot Water Boiler',
                fuel,
                f'{float(eff)*100:.0f}%' if eff else '80%',
                '[Capacity TBD]'
            ])
        add_equipment_table(doc, 'Hot Water Boilers', ['Name', 'Type', 'Fuel', 'Efficiency', 'Capacity'], boiler_data)

    # Water-to-Air Heat Pumps
    add_formatted_heading(doc, 'Water-to-Air Heat Pump Components', level=2)

    doc.add_paragraph(f'The model includes {len(cooling_coils_wshp)} water-to-air heat pump cooling coils and {len(heating_coils_wshp)} heating coils serving individual thermal zones.')
    doc.add_paragraph()

    wshp_sample = []
    for i, (cool, heat) in enumerate(zip(cooling_coils_wshp[:5], heating_coils_wshp[:5]), 1):
        wshp_sample.append([
            str(i),
            cool.Name[:50] + '...' if len(cool.Name) > 50 else cool.Name,
            heat.Name[:50] + '...' if len(heat.Name) > 50 else heat.Name
        ])

    if len(cooling_coils_wshp) > 5:
        wshp_sample.append(['...', f'[{len(cooling_coils_wshp) - 5} more cooling coils]', f'[{len(heating_coils_wshp) - 5} more heating coils]'])

    add_equipment_table(doc, 'Sample WSHP Coils (first 5)', ['#', 'Cooling Coil', 'Heating Coil'], wshp_sample)

    # Energy Recovery
    if ervs:
        add_formatted_heading(doc, 'Energy Recovery Ventilation', level=2)

        erv_data = []
        for erv in ervs:
            config = 'Rotary Wheel' if 'Rotary' in erv.Name or 'Wheel' in erv.Name else 'Plate'
            erv_data.append([
                erv.Name,
                'Air-to-Air Sensible and Latent',
                config,
                '[Effectiveness TBD]'
            ])

        add_equipment_table(doc, '', ['Name', 'Type', 'Configuration', 'Effectiveness'], erv_data)

    # Pumps
    if all_pumps:
        add_formatted_heading(doc, 'Pumping Systems', level=2)

        pump_data = []
        for pump in all_pumps:
            pump_type = 'Variable Speed' if pump.key.upper() == 'PUMP:VARIABLESPEED' else 'Constant Speed'
            power = getattr(pump, 'Rated_Power_Consumption', 'Autosized')
            flow = getattr(pump, 'Rated_Flow_Rate', 'Autosized')

            pump_data.append([
                pump.Name,
                pump_type,
                str(power),
                str(flow)
            ])

        add_equipment_table(doc, '', ['Pump Name', 'Type', 'Power (W)', 'Flow (m³/s)'], pump_data)

    # Placeholder sections for manual input
    doc.add_page_break()
    add_formatted_heading(doc, 'Modeling Results (To Be Completed)', level=1)

    doc.add_paragraph('[Insert Table 1: Modeling Results Summary here]')
    doc.add_paragraph()
    doc.add_paragraph('Package C results will be populated after baseline comparison.')
    doc.add_paragraph()

    # Table 1 placeholder structure
    results_table = doc.add_table(rows=2, cols=7)
    results_table.style = 'Light Grid Accent 1'

    headers = ['Design Package', 'PBPm ($/yr)', 'Energy Cost Savings ($/yr)', 'Percent Below Code (%)', 'Life Cycle NPV ($)', 'Peak Cooling (tons)', 'Peak Heating (kBtu/h)']
    header_cells = results_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Package C row (placeholders)
    package_c_cells = results_table.rows[1].cells
    package_c_cells[0].text = 'Package C: GSHP'
    package_c_cells[1].text = '[TBD]'
    package_c_cells[2].text = '[TBD]'
    package_c_cells[3].text = '[TBD]'
    package_c_cells[4].text = '[TBD]'
    package_c_cells[5].text = '[TBD]'
    package_c_cells[6].text = '[TBD]'

    doc.add_paragraph()

    # Add notes section
    add_formatted_heading(doc, 'Notes', level=2)

    doc.add_paragraph('1. Energy cost savings calculated as BBPCode - PBPm')
    doc.add_paragraph('2. Percent below code calculated as: 1 - PBREC/(BPFCode × BBREC)')
    doc.add_paragraph('3. Peak loads are coil loads from model, not final equipment sizes')
    doc.add_paragraph('4. All values pending baseline model completion')

    # Save document
    doc.save(output_path)

    print(f"\n[OK] IDAP addendum document created: {output_path}\n")
    print("Document sections:")
    print(f"  - Energy Model Input Details (HVAC systems)")
    print(f"  - Equipment Schedule:")
    print(f"    • {len(airloops)} Air Systems")
    print(f"    • {len(fluid_coolers)} Fluid Coolers")
    print(f"    • {len(boilers)} Boilers")
    print(f"    • {len(cooling_coils_wshp)} WSHP Cooling Coils")
    print(f"    • {len(heating_coils_wshp)} WSHP Heating Coils")
    print(f"    • {len(ervs)} ERV Units")
    print(f"    • {len(all_pumps)} Pumps")
    print(f"  - Modeling Results (placeholders for manual completion)")

if __name__ == "__main__":
    from datetime import datetime
    main()
